import re
import numpy
from docx import Document
import json


class Reconstructor:
    def __init__(self, file_name):
        self.file_name = file_name
        self.document = Document(self.file_name)
        self.groups_name = []
        self.tables = []
        self.couples = []

        self.run_sync_functions()

    def run_sync_functions(self):
        self.get_groups()
        self.grouping_a_tables()
        self.get_couples_of_groups()

    def get_groups(self):
        for paragraph in self.document.paragraphs:
            if re.search('/', paragraph.text) is None:
                if re.search('ЧИСЛИТЕЛЬ', paragraph.text):
                    self.groups_name.append(paragraph.text.split('ГРУППЫ')[1].split('неделя')[0].strip())
            else:
                if re.search('ЧИСЛИТЕЛЬ', paragraph.text):
                    self.groups_name.append(paragraph.text.split('ГРУППЫ')[1].split('неделя')[0].strip().split('/')[0])
                    self.groups_name.append(paragraph.text.split('ГРУППЫ')[1].split('неделя')[0].strip().split('/')[1])

    def get_couples_of_groups(self):
        for tables in self.tables:
            couples = self.make_a_blank()
            for counter_table, table in enumerate(tables):
                for row in table.columns:
                    if row.cells[0].text in couples['numerator'].keys():
                        for counter, cell in enumerate(row.cells):
                            if cell.text not in couples['numerator'].keys():
                                if counter_table == 0:
                                    couples['numerator'][row.cells[0].text][counter].append(str(cell.text).encode('utf-8').decode())
                                else:
                                    couples['denominator'][row.cells[0].text][counter].append(str(cell.text).encode('utf-8').decode())
            self.couples.append(couples)

    def grouping_a_tables(self):
        self.tables = numpy.array_split(self.document.tables, len(self.document.tables)//2)

    @staticmethod
    def make_a_blank() -> dict:
        blank = {'numerator': {'Понедельник': {1: ['09:20-10:55'], 2: ['11:05-12:40'], 3: ['13:20-14:55'], 4: ['15:05-16:40']},
                               'Вторник': {1: ['09:20-10:55'], 2: ['11:05-12:40'], 3: ['13:20-14:55'], 4: ['15:05-16:40']},
                               'Среда': {1: ['09:20-10:55'], 2: ['11:05-12:40'], 3: ['13:20-14:55'], 4: ['15:05-16:40']},
                               'Четверг': {1: ['09:20-10:55'], 2: ['11:05-12:40'], 3: ['13:20-14:55'], 4: ['15:05-16:40']},
                               'Пятница': {1: ['09:20-10:55'], 2: ['11:05-12:40'], 3: ['13:20-14:55'], 4: ['15:05-16:40']},
                               'Суббота': {1: ['09:20-10:55'], 2: ['11:05-12:40'], 3: ['13:20-14:55'], 4: ['15:05-16:40']}},
                 'denominator': {'Понедельник': {1: ['09:20-10:55'], 2: ['11:05-12:40'], 3: ['13:20-14:55'], 4: ['15:05-16:40']},
                                 'Вторник': {1: ['09:20-10:55'], 2: ['11:05-12:40'], 3: ['13:20-14:55'], 4: ['15:05-16:40']},
                                 'Среда': {1: ['09:20-10:55'], 2: ['11:05-12:40'], 3: ['13:20-14:55'], 4: ['15:05-16:40']},
                                 'Четверг': {1: ['09:20-10:55'], 2: ['11:05-12:40'], 3: ['13:20-14:55'], 4: ['15:05-16:40']},
                                 'Пятница': {1: ['09:20-10:55'], 2: ['11:05-12:40'], 3: ['13:20-14:55'], 4: ['15:05-16:40']},
                                 'Суббота': {1: ['09:20-10:55'], 2: ['11:05-12:40'], 3: ['13:20-14:55'], 4: ['15:05-16:40']}}}
        return blank

    def send_data_to_db(self, self_bot):
        counter: int = 0
        for group, couple in zip(self.groups_name, self.couples):
            if group == 'С301к':
                self_bot.cursor.executemany("INSERT INTO timetable VALUES (%s, %s)", [(group, json.dumps(self.couples[counter-1], indent=4))])
            self_bot.cursor.executemany("INSERT INTO timetable VALUES (%s, %s)", [(group, json.dumps(couple, indent=4))])
            counter += 1
        self_bot.connection.commit()
